# -*- coding: utf-8 -*-
"""
Génère le mémoire PFE complet en .docx conforme à la fiche technique.
- Times New Roman 12pt corps / 14pt gras titres bleu #1b3a6b
- Interligne 1,5 / justifié / A4
- Marges : Haut/Bas 2,5cm, Gauche 3cm, Droite 2,5cm
- Pagination bas centré
- Tableaux numérotés avec en-tête bleu, figures encadrées
"""
import os, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"C:\Users\PC\Desktop\elawael-platform\docs\memoire"
OUT  = r"C:\Users\PC\Desktop\elawael-platform\Rapport_PFE_Marzouk_ElAwael.docx"
FILES = ["00-pages-liminaires.md","01-introduction.md","02-partie1-cadre.md",
         "03-partie2-conception.md","04-partie3-realisation.md","05-conclusion.md",
         "06-bibliographie.md","07-annexes.md"]

BLUE = RGBColor(0x1b,0x3a,0x6b)
WHITE = RGBColor(0xFF,0xFF,0xFF)
FONT = "Times New Roman"
BOX = re.compile(r'[═║│─└┘┌┐├┤┬┴┼╔╗╚╝╠╣╦╩╬]')
ARABIC = re.compile(r'[؀-ۿ]')

def set_font(run, size=12, bold=False, italic=False, color=None, name=FONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    # assure la police pour caractères complexes/arabes
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii','w:hAnsi','w:cs'):
        rfonts.set(qn(a), name)

def add_inline(p, text, base_size=12, base_color=None, base_italic=False):
    """Gère le gras **...** dans une ligne."""
    if ARABIC.search(text):
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ppr = p._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi'); ppr.append(bidi)
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); set_font(r, base_size, True, base_italic, base_color)
        else:
            # italique *...* simple
            sub = re.split(r'(\*[^*]+?\*)', part)
            for s in sub:
                if not s: continue
                if s.startswith('*') and s.endswith('*') and len(s) > 2:
                    r = p.add_run(s[1:-1]); set_font(r, base_size, False, True, base_color)
                else:
                    r = p.add_run(s); set_font(r, base_size, False, base_italic, base_color)

def shade(cell, hexc):
    tcpr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),hexc)
    tcpr.append(sh)

def set_cell_borders(cell):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        e = OxmlElement('w:'+edge)
        e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4')
        e.set(qn('w:space'),'0'); e.set(qn('w:color'),'B0B8C4')
        borders.append(e)
    tcpr.append(borders)

doc = Document()

# ---- Style Normal ----
normal = doc.styles['Normal']
normal.font.name = FONT
normal.font.size = Pt(12)
normal.element.rPr.rFonts.set(qn('w:cs'), FONT)
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf.space_after = Pt(6)

# ---- Marges & A4 + pagination ----
sec = doc.sections[0]
sec.page_width = Cm(21); sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(3); sec.right_margin = Cm(2.5)

def add_page_number_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.text = ''
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'),'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve'); instr.text='PAGE'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'),'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
    set_font(run, 11)

add_page_number_footer(sec)

def heading(text, level):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 1:
        p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(12)
        size = 14; bold = True; color = BLUE
        p.paragraph_format.keep_with_next = True
    elif level == 2:
        p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
        size = 14; bold = True; color = BLUE
        p.paragraph_format.keep_with_next = True
    elif level == 3:
        p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
        size = 13; bold = True; color = BLUE
        p.paragraph_format.keep_with_next = True
    else:
        p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
        size = 12; bold = True; color = BLUE
    r = p.add_run(text); set_font(r, size, bold, False, color)
    # outline level pour la table des matières
    ppr = p._p.get_or_add_pPr()
    ol = OxmlElement('w:outlineLvl'); ol.set(qn('w:val'), str(level-1)); ppr.append(ol)
    return p

def normal_para(text):
    p = doc.add_paragraph()
    add_inline(p, text)
    return p

def quote_para(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    add_inline(p, text, base_italic=True, base_color=RGBColor(0x33,0x44,0x55))
    # bordure gauche
    ppr = p._p.get_or_add_pPr()
    pbd = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),'single'); left.set(qn('w:sz'),'18'); left.set(qn('w:space'),'8'); left.set(qn('w:color'),'1b3a6b')
    pbd.append(left); ppr.append(pbd)
    return p

def bullet(text, ordered=False):
    p = doc.add_paragraph(style='List Number' if ordered else 'List Bullet')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    add_inline(p, text)
    return p

def title_page(lines):
    for ln in lines:
        clean = BOX.sub('', ln).strip()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        if not clean:
            r = p.add_run(''); set_font(r, 6); continue
        up = clean.isupper() and len(clean) > 3
        big = any(k in clean for k in ['MÉMOIRE DE FIN','MASTER EN LEADERSHIP','CONCEPTION ET RÉALISATION'])
        r = p.add_run(clean)
        set_font(r, 16 if big else (13 if up else 12), big or up, False, BLUE if (big or up) else None)
    doc.add_page_break()

def figure_box(lines):
    """Encadré gris pour les emplacements de figures / blocs ASCII."""
    txt = [l for l in lines]
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0,0)
    shade(cell, 'F0F3F8'); set_cell_borders(cell)
    cell.paragraphs[0].text = ''
    first = True
    is_fig = txt and txt[0].strip().startswith('[FIGURE')
    for l in txt:
        cl = BOX.sub('', l).rstrip()
        if first:
            p = cell.paragraphs[0]; first = False
        else:
            p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(cl if cl else ' ')
        if is_fig and l.strip().startswith('[FIGURE'):
            set_font(r, 11, True, False, BLUE)
        elif l.strip().startswith('Source'):
            set_font(r, 9, False, True, RGBColor(0x66,0x66,0x66))
        else:
            set_font(r, 9.5, False, l.strip().startswith('Section'), RGBColor(0x44,0x44,0x44), name='Consolas')
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def render_table(rows):
    # rows: liste de listes de cellules ; rows[1] est le séparateur ---
    header = rows[0]
    body = rows[2:]
    ncol = len(header)
    tbl = doc.add_table(rows=1, cols=ncol)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True
    hdr = tbl.rows[0].cells
    for i,h in enumerate(header):
        shade(hdr[i],'1b3a6b'); set_cell_borders(hdr[i])
        para = hdr[i].paragraphs[0]; para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        para.paragraph_format.space_after = Pt(2)
        add_inline_cell(para, h.strip(), 11, True, WHITE)
    for ri,row in enumerate(body):
        cells = tbl.add_row().cells
        for i in range(ncol):
            val = row[i].strip() if i < len(row) else ''
            set_cell_borders(cells[i])
            if ri % 2 == 1: shade(cells[i],'F4F7FB')
            para = cells[i].paragraphs[0]
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            para.paragraph_format.space_after = Pt(2)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_cell(para, val, 10.5, False, None)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_inline_cell(p, text, size, bold, color):
    if ARABIC.search(text):
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if not part: continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); set_font(r, size, True, False, color)
        else:
            r = p.add_run(part); set_font(r, size, bold, False, color)

def parse_table_row(line):
    s = line.strip()
    if s.startswith('|'): s = s[1:]
    if s.endswith('|'): s = s[:-1]
    return s.split('|')

# ---------- Lecture + rendu ----------
for fi, fname in enumerate(FILES):
    path = os.path.join(BASE, fname)
    with open(path, encoding='utf-8') as f:
        lines = f.read().split('\n')
    if fi > 0:
        doc.add_page_break()
    i = 0
    in_code = False
    code_buf = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        # bloc code
        if stripped.startswith('```'):
            if not in_code:
                in_code = True; code_buf = []
            else:
                in_code = False
                # décider : page de garde ou figure/ascii
                joined = '\n'.join(code_buf)
                if 'MÉMOIRE DE FIN' in joined or 'MASTER EN LEADERSHIP' in joined:
                    title_page(code_buf)
                else:
                    figure_box(code_buf)
            i += 1; continue
        if in_code:
            code_buf.append(line); i += 1; continue
        # titres
        if stripped.startswith('#### '):
            heading(stripped[5:].strip(), 4); i += 1; continue
        if stripped.startswith('### '):
            heading(stripped[4:].strip(), 3); i += 1; continue
        if stripped.startswith('## '):
            heading(stripped[3:].strip(), 2); i += 1; continue
        if stripped.startswith('# '):
            heading(stripped[2:].strip(), 1); i += 1; continue
        # séparateur horizontal
        if stripped == '---':
            i += 1; continue
        # tableau
        if stripped.startswith('|') and i+1 < len(lines) and re.match(r'^\|?[\s:\-|]+\|?$', lines[i+1].strip()) and '-' in lines[i+1]:
            tbl_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl_rows.append(parse_table_row(lines[i])); i += 1
            render_table(tbl_rows); continue
        # citation
        if stripped.startswith('>'):
            quote_para(stripped.lstrip('>').strip()); i += 1; continue
        # listes
        m_ol = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if stripped.startswith('- '):
            bullet(stripped[2:].strip(), ordered=False); i += 1; continue
        if m_ol:
            bullet(m_ol.group(2).strip(), ordered=True); i += 1; continue
        # ligne vide
        if not stripped:
            i += 1; continue
        # paragraphe normal
        normal_para(stripped); i += 1

doc.save(OUT)
print("OK ->", OUT)
print("Paragraphes:", len(doc.paragraphs), "| Tableaux:", len(doc.tables))
